"""
Assembles report/Final_Report.docx from the artifacts produced by
Stochastic_Price_Simulation.ipynb (tables in output/, figures in figures/).

Run after the notebook has been executed at least once:
    .venv/Scripts/python.exe report/generate_report.py
"""

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
FIGURES = ROOT / "figures"
OUTPUT = ROOT / "output"
TICKER = "TSLA"

stats = pd.read_csv(OUTPUT / "summary_statistics.csv", index_col=0)
sim = pd.read_csv(OUTPUT / "simulation_summary.csv")
prices = pd.read_csv(OUTPUT / "clean_close_prices.csv", index_col=0, parse_dates=True)
backtest = pd.read_csv(OUTPUT / "backtest_results.csv", index_col=0)

NAVY = RGBColor(0x1F, 0x2D, 0x50)
GRAY = RGBColor(0x55, 0x55, 0x55)

doc = Document()

# ---- base styles -----------------------------------------------------------
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)

for style_name, size, bold in [("Heading 1", 18, True), ("Heading 2", 14, True), ("Heading 3", 12, True)]:
    st = doc.styles[style_name]
    st.font.name = "Calibri"
    st.font.size = Pt(size)
    st.font.bold = bold
    st.font.color.rgb = NAVY


def caption(text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = GRAY


def figure(path: Path, cap: str, width: float = 6.3):
    doc.add_picture(str(path), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption(cap)


def add_table(df: pd.DataFrame, header_bold: bool = True):
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr[i].text = str(col)
        if header_bold:
            hdr[i].paragraphs[0].runs[0].font.bold = True
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = f"{val:.4f}" if isinstance(val, float) else str(val)
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9.5)
    return table


# =========================================================================== #
# Title page
# =========================================================================== #
title = doc.add_heading("Stochastic Processes Problem Set:", level=0)
title.runs[0].font.color.rgb = NAVY
sub = doc.add_heading(
    "Modeling Tesla (TSLA) Stock Price Movements with GBM, Ornstein-Uhlenbeck, "
    "and Merton Jump-Diffusion Processes", level=1)
sub.runs[0].font.color.rgb = NAVY

meta = doc.add_paragraph()
meta.add_run("\nAuthor: ").bold = True
meta.add_run("Alberto Kabore\n")
meta.add_run("Role: ").bold = True
meta.add_run("Quantitative Analyst, Financial Research Team (course use case)\n")
meta.add_run("Course: ").bold = True
meta.add_run("Stochastic Processes — Assignment 2\n")
meta.add_run("Date: ").bold = True
meta.add_run("August 26, 2026\n")
doc.add_page_break()

# =========================================================================== #
# 1. Introduction
# =========================================================================== #
doc.add_heading("1. Introduction and Objective", level=1)
doc.add_paragraph(
    "This report applies stochastic process theory to the practical problem faced by a "
    "quantitative analyst on a financial research team: characterizing the future distribution "
    "of a stock's price so that investors can reason quantitatively about risk and return under "
    "uncertainty. Three continuous-time stochastic models are calibrated to real historical price "
    "data for Tesla, Inc. (TSLA) and used to run Monte Carlo simulations: Geometric Brownian "
    "Motion (GBM), the Ornstein-Uhlenbeck (OU) mean-reverting process, and a Merton "
    "Jump-Diffusion (MJD) process. The objective is not only to simulate plausible future price "
    "paths, but to compare how each model's structural assumptions shape its risk output, and to "
    "translate that comparison into concrete guidance for investors."
)
doc.add_paragraph(
    "Tesla was selected because its return history — hyper-growth, multiple stock splits, sharp "
    "drawdowns, and large single-day, news-driven moves — exercises all three models' assumptions "
    "under genuinely turbulent conditions, making it a strong single-name test case for comparing "
    "a trend model (GBM), a mean-reversion model (OU), and a fat-tail model (MJD). All code, "
    "calibration, simulation, and visualization logic accompanies this report as a Jupyter "
    "notebook (Stochastic_Price_Simulation.ipynb) and a reusable Python module "
    "(src/stochastic_models.py)."
)

# =========================================================================== #
# 2. Data & preprocessing
# =========================================================================== #
doc.add_heading("2. Description of Data and Preprocessing", level=1)
doc.add_paragraph(
    f"Daily OHLCV data for Tesla was sourced from the Kaggle dataset "
    f"“jillanisofttech/tesla-stock-price” (Jillani SoftTech, 2024), downloaded via the "
    f"kagglehub API and used as the sole data source for this project. The dataset's longer file "
    f"(Tasla_Stock_Updated_V2.csv) was used, spanning {prices.index.min().date()} through "
    f"{prices.index.max().date()} — {stats.loc[TICKER, 'n_obs']:.0f} trading days after cleaning, "
    f"far exceeding the one-year minimum. Prices are already split-adjusted, so day-over-day log "
    f"returns reflect true economic price changes rather than stock-split artifacts."
)
doc.add_heading("2.1 Cleaning pipeline", level=2)
doc.add_paragraph(
    "The raw CSV was passed through a deterministic cleaning routine before any modeling: (1) the "
    "series was reindexed to a business-day calendar and short gaps were forward-filled; (2) rows "
    "with missing, non-finite, or non-positive close prices were dropped rather than imputed, to "
    "avoid fabricating data; (3) large single-day return outliers were deliberately retained, since "
    "they are genuine market-moving events (e.g., earnings surprises, delivery-number surprises) "
    "that the jump-diffusion model is specifically designed to capture — removing them would "
    "understate tail risk in every model."
)
doc.add_heading("2.2 Exploratory statistics", level=2)
doc.add_paragraph(
    "Table 1 summarizes annualized return and volatility, return-distribution shape (skewness, "
    "excess kurtosis), and Augmented Dickey-Fuller (ADF) stationarity test p-values for both the "
    "price level and the log-return series."
)
tbl1 = stats.loc[[TICKER], ["n_obs", "last_price", "ann_mean_return", "ann_volatility", "skew",
                             "excess_kurtosis", "adf_pvalue_price", "adf_pvalue_returns"]].reset_index()
tbl1.columns = ["Ticker", "N (days)", "Last Price", "Ann. Return", "Ann. Volatility",
                "Skew", "Excess Kurtosis", "ADF p (price)", "ADF p (returns)"]
add_table(tbl1)
caption(f"Table 1. Descriptive and stationarity statistics, {prices.index.min().date()} – "
        f"{prices.index.max().date()}.")
doc.add_paragraph(
    "The price-level ADF test fails to reject a unit root (p > 0.5), consistent with price "
    "following a random walk, while the return series is strongly stationary (p ≈ 0) — exactly "
    "the pattern GBM assumes and OU (applied to price levels) does not. Excess kurtosis is well "
    "above 0, indicating fat-tailed return behavior that directly motivates the jump-diffusion "
    "model in Section 3.3."
)
figure(FIGURES / "fig01_price_history.png",
       "Figure 1. TSLA daily closing price, 2015–2024 (Kaggle dataset).")

doc.add_heading("2.3 Formal statistical tests", level=2)
doc.add_paragraph(
    "Beyond the descriptive skew/kurtosis figures in Table 1, two formal hypothesis tests — "
    "computed with SciPy and statsmodels, per the assignment's suggested toolset — test the "
    "models' core assumptions directly rather than relying on visual inspection."
)
jb_stat, jb_p = stats.loc[TICKER, "jarque_bera_stat"], stats.loc[TICKER, "jarque_bera_pvalue"]
lb_stat, lb_p = stats.loc[TICKER, "ljung_box_stat_lag10"], stats.loc[TICKER, "ljung_box_pvalue_lag10"]
tbl_tests = pd.DataFrame({
    "Test": ["Jarque-Bera (scipy.stats.jarque_bera)",
             "Ljung-Box on squared returns, lag 10 (statsmodels)"],
    "Null hypothesis": ["Returns are Normally distributed", "No autocorrelation in squared returns"],
    "Statistic": [f"{jb_stat:.1f}", f"{lb_stat:.1f}"],
    "p-value": [f"{jb_p:.2e}", f"{lb_p:.2e}"],
    "Verdict": ["Rejected" if jb_p < 0.05 else "Not rejected",
                "Rejected" if lb_p < 0.05 else "Not rejected"],
})
add_table(tbl_tests)
caption("Table 2. Formal hypothesis tests on TSLA daily log returns.")
doc.add_paragraph(
    "The Jarque-Bera test rejects normality of returns at the 5% level — direct statistical "
    "confirmation (not just a descriptive kurtosis number) that GBM's i.i.d. Normal-return "
    "assumption does not hold for this stock, motivating the jump-diffusion model in Section "
    "3.3. The Ljung-Box test on squared returns rejects the no-autocorrelation null, evidence of "
    "volatility clustering — large price moves tend to be followed by more large moves. None of "
    "the three models in this study can reproduce that pattern, since each draws independent "
    "random increments at every simulated step; this is revisited as a shared limitation in "
    "Section 7.1."
)

# =========================================================================== #
# 3. Models
# =========================================================================== #
doc.add_heading("3. Explanation of Models and Implementation", level=1)
doc.add_paragraph(
    "Two models are the minimum required for comparison; three are implemented here to give a "
    "fuller picture of the trade-offs involved, and specifically to include one model (OU) whose "
    "assumptions are deliberately mismatched to a structurally growing equity, as a pedagogical "
    "contrast."
)

doc.add_heading("3.1 Geometric Brownian Motion (GBM)", level=2)
doc.add_paragraph(
    "GBM is the standard model underlying Black-Scholes-Merton option pricing (Black & Scholes, "
    "1973): dS_t = μS_t dt + σS_t dW_t, with exact solution "
    "S_t = S_0 exp[(μ − σ²/2)t + σW_t]. Annualized drift μ and volatility σ are estimated "
    "directly from the sample mean and standard deviation of daily log returns and simulated using "
    "the exact log-normal solution (no discretization bias)."
)
doc.add_heading("3.2 Ornstein-Uhlenbeck (OU) process", level=2)
doc.add_paragraph(
    "The OU process (Uhlenbeck & Ornstein, 1930), dX_t = θ(μ − X_t)dt + σdW_t, pulls the level "
    "back toward a long-run mean μ at speed θ. It was calibrated by recognizing that its exact "
    "discretization is a Gaussian AR(1) model, fitting that regression by ordinary least squares, "
    "and analytically recovering θ, μ, and σ. It is included specifically as a contrast case: "
    "Tesla's price is not economically expected to mean-revert to a fixed level — the stock grew "
    "by more than an order of magnitude over the sample window — and Section 2's ADF results "
    "support that mismatch; it is examined directly in Section 5."
)
doc.add_heading("3.3 Merton Jump-Diffusion (MJD)", level=2)
doc.add_paragraph(
    "The Merton (1976) jump-diffusion model augments GBM with a compound Poisson jump process: "
    "dS_t/S_t = (μ − λk)dt + σdW_t + dJ_t, where jumps arrive at annual intensity λ and log-jump "
    "sizes are Normal(μ_J, σ_J²); k = E[Y−1] compensates the drift so jumps do not, by "
    "themselves, bias the mean. Calibration flags daily log returns more than three standard "
    "deviations from the mean as jump days, fits the jump distribution from those observations, "
    "and fits the diffusive μ, σ from the remaining returns."
)

doc.add_paragraph(
    "Full calibrated parameters for all three models are reported in the accompanying notebook "
    "(Sections 4–6) and in output/summary_statistics.csv; they are referenced throughout the "
    "interpretation below rather than reproduced in full here."
)

# =========================================================================== #
# 4. Model validation
# =========================================================================== #
doc.add_heading("4. Model Validation", level=1)
doc.add_paragraph(
    "Before using the three calibrated models for forward-looking simulation, two independent "
    "checks confirm the simulators are implemented correctly and are not badly miscalibrated "
    "against data they were not fit on."
)
doc.add_heading("4.1 Analytic cross-check (GBM)", level=2)
doc.add_paragraph(
    "GBM has a closed-form log-normal terminal distribution, so its Monte Carlo output can be "
    "checked against exact theory rather than only against itself. The 5th/50th/95th percentiles "
    "computed analytically via scipy.stats.norm.ppf from the log-normal density agree with the "
    "empirical percentiles from a 20,000-path Monte Carlo run to within roughly 1%, confirming "
    "the simulator has no implementation bias (full comparison table in the accompanying "
    "notebook, Section 7.1)."
)
doc.add_heading("4.2 Out-of-sample backtest (all three models)", level=2)
doc.add_paragraph(
    "Each model was refit on data with the last 180 trading days withheld, simulated 180 days "
    "forward from that earlier point, and compared against the price actually realized over that "
    "period."
)
bt_display = backtest.reset_index().rename(columns={"index": "Model"})
bt_display = bt_display[["Model", "S0", "actual_terminal", "sim_mean_terminal",
                          "sim_p05", "sim_p95", "percentile_rank_of_actual", "inside_90pct_interval"]]
bt_display.columns = ["Model", "S0", "Actual Terminal", "Sim. Mean Terminal",
                       "Sim. p05", "Sim. p95", "Percentile Rank", "Inside 90% CI"]
add_table(bt_display)
caption("Table 3. Out-of-sample backtest: 180-day holdout, trained through the date 180 trading "
        "days before the end of the dataset.")
n_inside = int(backtest["inside_90pct_interval"].sum())
doc.add_paragraph(
    f"The actual realized TSLA price over the holdout period landed inside all {n_inside} of the "
    f"three models' own 90% confidence intervals for this particular window — a reassuring but "
    f"limited result, since it reflects a single historical episode rather than a rolling, "
    f"multi-window validation. OU's percentile rank is noticeably higher than GBM's or MJD's "
    f"because OU's fitted mean reversion predicted a pullback that did not occur; TSLA continued "
    f"climbing over the holdout period, illustrating in a concrete, out-of-sample way exactly the "
    f"mismatch discussed qualitatively in Section 3.2. Extending this into a rolling walk-forward "
    f"backtest is recommended before any live use (Section 7.2, Recommendation 6)."
)

# =========================================================================== #
# 5. Simulation methodology
# =========================================================================== #
doc.add_heading("5. Monte Carlo Simulation Methodology", level=1)
doc.add_paragraph(
    "For each of the three models and three horizons (30, 60, and 180 trading days — roughly "
    "six weeks, three months, and nine months), 1,000 independent price paths were simulated from "
    "the last observed close in the dataset. A common random-number seed was reused per horizon "
    "across models so the comparison is as apples-to-apples as each model's structure permits."
)

# =========================================================================== #
# 6. Results & interpretation
# =========================================================================== #
doc.add_heading("6. Simulation Results and Interpretation", level=1)

figure(FIGURES / "fig02_paths.png",
       "Figure 2. TSLA: 1,000-path Monte Carlo simulation, 180-day horizon, by model. Shaded "
       "band is the 5th–95th percentile envelope.")
figure(FIGURES / "fig03_terminal_hist.png",
       "Figure 3. TSLA: terminal price distribution after 180 trading days.")

row180 = sim.query("horizon_days==180").set_index("model")
gbm, ou, mjd = row180.loc["GBM"], row180.loc["OU"], row180.loc["MJD"]
doc.add_paragraph(
    f"At the 180-trading-day horizon, GBM projects a mean terminal price of "
    f"${gbm['mean_terminal']:.2f} with a {gbm['prob_loss_pct']:.1f}% probability of loss and a "
    f"95% Value-at-Risk (VaR) of {gbm['VaR_95_pct']:.1f}%. MJD — which additionally prices in "
    f"jump risk — reports a {mjd['prob_loss_pct']:.1f}% probability of loss and a "
    f"{mjd['VaR_95_pct']:.1f}% 95% VaR (95% CVaR {mjd['CVaR_95_pct']:.1f}%). OU reports a "
    f"{ou['prob_loss_pct']:.1f}% probability of loss with a visibly narrower confidence band in "
    f"Figure 2, reflecting its bounded long-run variance rather than GBM/MJD's variance that "
    f"grows with the horizon."
)

doc.add_heading("6.1 Cross-model risk comparison", level=2)
figure(FIGURES / "fig04_risk_comparison.png",
       "Figure 4. Probability of loss, 95% VaR, and return-to-VaR ratio by model and horizon.")
doc.add_paragraph(
    "MJD reports equal or higher tail risk (probability of loss and VaR) than GBM at every "
    "horizon, consistent with MJD nesting GBM's diffusive risk and adding jump risk on top. Given "
    "Tesla's elevated excess kurtosis (Table 1), its statistically confirmed non-Normal returns "
    "(Table 2), and its well-documented history of large single-day, news-driven moves, this gap "
    "between GBM and MJD is one of the most actionable results in this study: GBM alone "
    "understates the odds of a large adverse move for this stock. OU's confidence bands are "
    "visibly narrower at the 180-day horizon, a direct consequence of mean reversion capping the "
    "long-run variance — a structural property, not evidence that OU is the more accurate model "
    "for this instrument (see Section 7)."
)

doc.add_heading("6.2 Short-term vs. long-term behavior", level=2)
doc.add_paragraph(
    "Under GBM and MJD, the standard deviation of the simulated terminal price grows roughly with "
    "the square root of the horizon, so a six-fold increase in horizon (30 → 180 days) does not "
    "produce a six-fold increase in terminal spread — uncertainty widens quickly at first and "
    "then more slowly. OU behaves qualitatively differently: because deviations from its fitted "
    "mean are pulled back at rate θ, its unconditional variance is bounded, so its confidence band "
    "stops widening noticeably beyond a horizon on the order of the process's half-life "
    "(ln 2 / θ, reported in the notebook). This distinction is the clearest evidence in the "
    "simulation output of how differently the three models extrapolate risk over long horizons."
)

doc.add_heading("6.3 Risk-adjusted comparison", level=2)
best_ratio_model = row180["return_to_VaR_ratio"].idxmax()
doc.add_paragraph(
    "The third panel of Figure 4 reports a simple risk-adjusted metric, the return-to-VaR ratio "
    "(mean simulated return divided by 95% VaR): expected return per unit of downside tail risk, "
    "which lets the three models be ranked on a single scale rather than comparing return and "
    f"risk separately. At the 180-day horizon, {best_ratio_model} ranks highest on this metric. "
    "This ratio is a useful summary statistic but not a substitute for the assumption checks in "
    "Sections 2.3 and 4 — a high risk-adjusted return is only meaningful if the underlying risk "
    "estimate is itself well-founded, and OU's apparently favorable risk figures at some horizons "
    "reflect its structurally narrower (mean-reverting) confidence band rather than a genuinely "
    "lower-risk forecast for this stock."
)

# =========================================================================== #
# 7. Realism, limitations, recommendations
# =========================================================================== #
doc.add_heading("7. Realism, Limitations, and Investment Recommendations", level=1)

doc.add_heading("7.1 Model realism and limitations", level=2)
for label, text in [
    ("GBM", "assumes constant volatility and i.i.d. Normally distributed log returns. The positive "
            "excess kurtosis measured for TSLA (Table 1) shows real returns have fatter tails than "
            "GBM allows, so GBM systematically understates the probability of extreme moves."),
    ("OU", "applied here to the raw price level, assumes price reverts to a fixed long-run mean — "
           "an assumption with no sound economic basis for a company that grew from roughly $14 to "
           "well over $200 a share over the sample window. The ADF test on the price level (Table "
           "1) fails to reject a unit root, directly contradicting OU's core stationarity "
           "assumption for this stock. It is far better suited to spreads, interest rates, or "
           "volatility levels — instruments with genuine economic anchors — than to a single, "
           "structurally growing equity."),
    ("MJD", "is the most complete of the three models but is sensitive to the jump-detection "
            "threshold used in calibration (3 standard deviations here) and assumes jump sizes are "
            "themselves Normally distributed and independent over time. The Ljung-Box test (Table "
            "2) confirms real market shocks cluster in time (volatility clustering) — a dynamic "
            "none of these three memoryless models captures, since each draws independent random "
            "increments at every simulated step."),
]:
    p = doc.add_paragraph()
    p.add_run(f"{label}: ").bold = True
    p.add_run(text)

doc.add_paragraph(
    "All three models hold their parameters fixed for the entire simulation horizon and ignore "
    "regime shifts, changing macroeconomic conditions, and correlation with other assets — a "
    "meaningful gap for horizons as long as 180 trading days, and especially relevant here since "
    "the underlying dataset is a static historical snapshot ending January 16, 2024, and is not "
    "being refreshed with more recent price action."
)

doc.add_heading("7.2 Recommendations for investors", level=2)
recs = [
    "Use GBM as a baseline for central-tendency planning (expected price, typical range), but not "
    "as the sole basis for tail-risk or stress-testing decisions on a name with TSLA's measured "
    "excess kurtosis and formally rejected return normality (Table 2), since it understates the "
    "odds of a large adverse move.",
    "Weight Merton Jump-Diffusion more heavily when sizing positions or setting stop-losses — its "
    "VaR/CVaR figures are the more conservative, and empirically better-supported, risk estimate "
    "for a jump-prone stock like Tesla; use the return-to-VaR ratio (Section 6.3) as a concrete, "
    "single-number way to weigh that trade-off against GBM's typically higher raw expected return.",
    "Do not use an outright-price OU model for directional forecasting of this stock; reserve it "
    "for spread trades, statistical arbitrage, or modeling mean-reverting quantities such as "
    "interest rates or volatility, where its assumptions are economically justified (Vasicek, "
    "1977).",
    "Treat longer-horizon forecasts with proportionally more skepticism, since all three models "
    "hold parameters fixed for the full horizon and are progressively more likely to be violated "
    "by real regime changes the further out the forecast extends.",
    "Recalibrate on current data before acting on this analysis in practice — the underlying "
    "dataset ends January 2024, and any real trading decision should refit these models on more "
    "recent prices rather than relying on parameters estimated from a static historical snapshot.",
    "Extend the single-window backtest in Section 4.2 into a rolling walk-forward validation "
    "before deploying any of these models in a live risk-management workflow — repeating the same "
    "holdout test across many overlapping historical windows would show whether the \"inside the "
    "90% interval\" result found here holds up consistently, or was specific to that one holdout "
    "period.",
]
for r in recs:
    doc.add_paragraph(r, style="List Number")

# =========================================================================== #
# 8. Conclusion
# =========================================================================== #
doc.add_heading("8. Conclusion", level=1)
doc.add_paragraph(
    "Three calibrated stochastic models — Geometric Brownian Motion, the Ornstein-Uhlenbeck "
    "process, and Merton Jump-Diffusion — were fit to nine years of daily closing prices for "
    "Tesla, Inc. (TSLA), sourced entirely from the Kaggle dataset "
    "“jillanisofttech/tesla-stock-price,” and used to run 1,000-path Monte Carlo simulations at "
    "30-, 60-, and 180-trading-day horizons. Beyond descriptive comparison, formal statistical "
    "tests (Jarque-Bera, Ljung-Box), an analytic cross-check of the simulator against closed-form "
    "GBM theory, and an out-of-sample backtest were used to validate both the implementation and "
    "the models' assumptions rather than relying on visual inspection alone. GBM provides a "
    "reasonable, simple baseline whose Normal-return assumption is formally rejected by the data; "
    "OU serves as an instructive theoretical contrast that highlights why mean-reversion models "
    "are poorly suited to a structurally growing equity; and Merton Jump-Diffusion offers the most "
    "complete picture of the fat-tailed, jump-prone behavior actually observed in Tesla's return "
    "history — though even MJD cannot reproduce the volatility clustering confirmed by the "
    "Ljung-Box test. All three share the limitation of fixed parameters over the simulation window "
    "and reliance on a static historical dataset, which argues for shorter forecast horizons, a "
    "rolling rather than single-window backtest, and periodic recalibration on refreshed data in "
    "any real trading or risk-management application."
)

# =========================================================================== #
# References
# =========================================================================== #
doc.add_heading("References", level=1)
refs = [
    "Black, F., & Scholes, M. (1973). The pricing of options and corporate liabilities. Journal of "
    "Political Economy, 81(3), 637–654. https://doi.org/10.1086/260062",
    "Cont, R., & Tankov, P. (2004). Financial modelling with jump processes. Chapman & Hall/CRC.",
    "Dickey, D. A., & Fuller, W. A. (1979). Distribution of the estimators for autoregressive time "
    "series with a unit root. Journal of the American Statistical Association, 74(366a), 427–431. "
    "https://doi.org/10.1080/01621459.1979.10482531",
    "Hull, J. C. (2022). Options, futures, and other derivatives (11th ed.). Pearson.",
    "Jarque, C. M., & Bera, A. K. (1980). Efficient tests for normality, homoscedasticity and "
    "serial independence of regression residuals. Economics Letters, 6(3), 255–259. "
    "https://doi.org/10.1016/0165-1765(80)90024-5",
    "Jillani SoftTech. (2024). Tesla stock price [Data set]. Kaggle. "
    "https://www.kaggle.com/datasets/jillanisofttech/tesla-stock-price",
    "Ljung, G. M., & Box, G. E. P. (1978). On a measure of lack of fit in time series models. "
    "Biometrika, 65(2), 297–303. https://doi.org/10.1093/biomet/65.2.297",
    "Merton, R. C. (1976). Option pricing when underlying stock returns are discontinuous. Journal "
    "of Financial Economics, 3(1–2), 125–144. https://doi.org/10.1016/0304-405X(76)90022-2",
    "Uhlenbeck, G. E., & Ornstein, L. S. (1930). On the theory of the Brownian motion. Physical "
    "Review, 36(5), 823–841. https://doi.org/10.1103/PhysRev.36.823",
    "Vasicek, O. (1977). An equilibrium characterization of the term structure. Journal of "
    "Financial Economics, 5(2), 177–188. https://doi.org/10.1016/0304-405X(77)90016-2",
]
for r in refs:
    p = doc.add_paragraph(r)
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)

out_path = ROOT / "report" / "Final_Report.docx"
doc.save(out_path)
print(f"wrote {out_path}")
